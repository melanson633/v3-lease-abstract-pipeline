"""I/O utilities for document loading and JSON serialization."""

from __future__ import annotations

import json
import re
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader

from leasegpt.errors import InputValidationError
from leasegpt.models import LeaseDocument

TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".json",
    ".csv",
    ".tsv",
    ".xml",
    ".html",
    ".htm",
    ".yaml",
    ".yml",
}


def utc_timestamp() -> str:
    return datetime.now(UTC).isoformat().replace("+00:00", "Z")


def load_json(path: Path) -> dict[str, Any]:
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        raise InputValidationError(f"Invalid JSON at {path}: {exc}") from exc


def write_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=True) + "\n", encoding="utf-8")


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def _parse_doc_code(path: Path) -> tuple[str, int]:
    stem_upper = path.stem.upper()
    if re.search(r"(^|[^A-Z0-9])OL([^A-Z0-9]|$)", stem_upper) or stem_upper.startswith("OL"):
        return "OL", 0
    if re.search(r"(^|[^A-Z0-9])CM([^A-Z0-9]|$)", stem_upper) or stem_upper.startswith("CM"):
        return "CM", 1
    amendment_match = re.search(r"(^|[^A-Z0-9])A(\d+)([^A-Z0-9]|$)", stem_upper)
    if amendment_match:
        number = int(amendment_match.group(2))
        return f"A{number}", 2_000 + number
    if re.search(r"(^|[^A-Z0-9])SL(\d+)([^A-Z0-9]|$)", stem_upper):
        sl_num = int(re.search(r"SL(\d+)", stem_upper).group(1))  # type: ignore[union-attr]
        return f"SL{sl_num}", 4_000 + sl_num
    return f"OTHER:{path.stem}", 9_000


def _read_pdf(path: Path) -> tuple[str, int]:
    reader = PdfReader(str(path))
    page_texts: list[str] = []
    for idx, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = text.strip()
        page_texts.append(f"[PAGE {idx}]\n{text}")
    joined = "\n\n".join(page_texts)
    return joined, len(reader.pages)


def _read_docx(path: Path) -> str:
    doc = Document(path)
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n".join(lines)


def _read_plain_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="latin-1")


def read_document_text(path: Path) -> tuple[str, int | None]:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _read_pdf(path)
    if ext == ".docx":
        return _read_docx(path), None
    if ext in TEXT_EXTENSIONS:
        return _read_plain_text(path), None
    raise InputValidationError(
        f"Unsupported file type '{ext}' for {path}. Supported: PDF, DOCX, and text-like files ({', '.join(sorted(TEXT_EXTENSIONS))})."
    )


def normalize_documents(paths: list[Path], max_chars_per_doc: int) -> list[LeaseDocument]:
    if not paths:
        raise InputValidationError("No input documents provided.")
    normalized: list[LeaseDocument] = []
    for path in paths:
        if not path.exists():
            raise InputValidationError(f"Input document not found: {path}")
        doc_code, order_index = _parse_doc_code(path)
        text, page_count = read_document_text(path)
        truncated = False
        trunc_note = None
        if len(text) > max_chars_per_doc:
            truncated = True
            text = text[:max_chars_per_doc]
            trunc_note = (
                f"Document text truncated at {max_chars_per_doc} characters for prompt safety."
            )
        normalized.append(
            LeaseDocument(
                path=path.resolve(),
                doc_code=doc_code,
                order_index=order_index,
                text=text,
                page_count=page_count,
                truncated=truncated,
                truncation_note=trunc_note,
            )
        )
    normalized.sort(key=lambda d: (d.order_index, d.path.name.lower()))
    return normalized


def coerce_bundle_shape(payload: dict[str, Any]) -> dict[str, Any]:
    if "lease_state" in payload:
        return {
            "lease_state": payload.get("lease_state") or {},
            "change_log": payload.get("change_log") or [],
            "pending_fields": payload.get("pending_fields") or [],
            "traceability": payload.get("traceability") or {"extractedFieldsMetadata": {}},
        }

    # Accept lease_state-only JSON.
    return {
        "lease_state": payload,
        "change_log": [],
        "pending_fields": [],
        "traceability": {"extractedFieldsMetadata": {}},
    }
